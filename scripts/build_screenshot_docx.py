from __future__ import annotations

from datetime import datetime
from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image


WORKSPACE = Path(r"C:\Users\user\Documents\New project")
OUT_DIR = WORKSPACE / "outputs" / "print"
ASSET_DIR = OUT_DIR / "docx-assets"
OUT_DOCX = OUT_DIR / "buildus-care-web-app-screenshots-2026-05-18.docx"

RAW_PATHS = [
    Path(r"C:\Users\user\Downloads\screencapture-buildus-care-flow-vercel-app-cases-2026-05-18-14_41_19.png"),
    Path(r"C:\Users\user\Downloads\screencapture-buildus-care-flow-vercel-app-request-photo-2026-05-18-14_41_30.png"),
    Path(r"C:\Users\user\Downloads\screencapture-buildus-care-flow-vercel-app-request-photo-2026-05-18-14_41_40.png"),
    Path(r"C:\Users\user\Downloads\screencapture-buildus-care-flow-vercel-app-request-photo-2026-05-18-14_41_57.png"),
    Path(r"C:\Users\user\Downloads\screencapture-buildus-care-flow-vercel-app-request-photo-2026-05-18-14_42_05.png"),
    Path(r"C:\Users\user\Downloads\screencapture-buildus-care-flow-vercel-app-request-photo-2026-05-18-14_42_41.png"),
    Path(r"C:\Users\user\Downloads\screencapture-buildus-care-flow-vercel-app-orders-lookup-2026-05-18-14_42_52.png"),
    Path(r"C:\Users\user\Downloads\screencapture-buildus-care-flow-vercel-app-orders-lookup-2026-05-18-14_43_23.png"),
    Path(r"C:\Users\user\Downloads\screencapture-buildus-care-flow-vercel-app-2026-05-18-14_34_59.png"),
    Path(r"C:\Users\user\Downloads\screencapture-buildus-care-flow-vercel-app-services-2026-05-18-14_35_20.png"),
    Path(r"C:\Users\user\Downloads\screencapture-buildus-care-flow-vercel-app-cases-2026-05-18-14_35_33.png"),
    Path(r"C:\Users\user\Downloads\screencapture-buildus-care-flow-vercel-app-request-photo-2026-05-18-14_35_44.png"),
    Path(r"C:\Users\user\Downloads\screencapture-buildus-care-flow-vercel-app-request-photo-2026-05-18-14_35_52.png"),
    Path(r"C:\Users\user\Downloads\screencapture-buildus-care-flow-vercel-app-request-photo-2026-05-18-14_36_09.png"),
    Path(r"C:\Users\user\Downloads\screencapture-buildus-care-flow-vercel-app-request-photo-2026-05-18-14_36_24.png"),
    Path(r"C:\Users\user\Downloads\screencapture-buildus-care-flow-vercel-app-request-photo-2026-05-18-14_36_32.png"),
    Path(r"C:\Users\user\Downloads\screencapture-buildus-care-flow-vercel-app-request-photo-2026-05-18-14_36_58.png"),
    Path(r"C:\Users\user\Downloads\screencapture-buildus-care-flow-vercel-app-orders-lookup-2026-05-18-14_37_12.png"),
    Path(r"C:\Users\user\Downloads\screencapture-buildus-care-flow-vercel-app-orders-lookup-2026-05-18-14_37_53.png"),
    Path(r"C:\Users\user\Downloads\screencapture-buildus-care-flow-vercel-app-quote-toilet-replace-2026-05-18-14_38_46.png"),
    Path(r"C:\Users\user\Downloads\screencapture-buildus-care-flow-vercel-app-quote-toilet-replace-2026-05-18-14_39_41.png"),
    Path(r"C:\Users\user\Downloads\screencapture-buildus-care-flow-vercel-app-orders-b369f71a-b304-481d-853c-73fa0938222e-2026-05-18-14_40_03.png"),
    Path(r"C:\Users\user\Downloads\screencapture-buildus-care-flow-vercel-app-orders-b369f71a-b304-481d-853c-73fa0938222e-2026-05-18-14_40_23.png"),
    Path(r"C:\Users\user\Downloads\screencapture-buildus-care-flow-vercel-app-quote-toilet-replace-2026-05-18-14_40_44.png"),
    Path(r"C:\Users\user\Downloads\screencapture-buildus-care-flow-vercel-app-2026-05-18-14_40_58.png"),
    Path(r"C:\Users\user\Downloads\screencapture-buildus-care-flow-vercel-app-services-2026-05-18-14_41_09.png"),
]

LABELS = {
    "services": "서비스 목록",
    "cases": "시공 사례",
    "request-photo": "사진 판정",
    "orders-lookup": "주문 조회",
    "quote-toilet-replace": "견적/결제",
}


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    tc_pr.append(shading)


def set_cell_border(cell, color: str = "E4E1DB") -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "8")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def parse_time(path: Path) -> datetime:
    match = re.search(r"(2026-05-18)-(\d{2})_(\d{2})_(\d{2})", path.name)
    if not match:
        return datetime.min
    return datetime.strptime(f"{match.group(1)} {match.group(2)}:{match.group(3)}:{match.group(4)}", "%Y-%m-%d %H:%M:%S")


def page_label(path: Path) -> str:
    name = path.name
    if re.search(r"app-2026-05-18", name):
        return "홈"
    if "orders-b369f71a" in name:
        return "주문 상세"
    for key, value in LABELS.items():
        if key in name:
            return value
    return "화면 캡처"


def classify(path: Path) -> str:
    with Image.open(path) as image:
        return "app" if image.width <= 800 else "web"


def set_section(section, orientation: str) -> None:
    if orientation == "landscape":
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = Inches(11.69)
        section.page_height = Inches(8.27)
        section.left_margin = Inches(0.35)
        section.right_margin = Inches(0.35)
        section.top_margin = Inches(0.35)
        section.bottom_margin = Inches(0.35)
    else:
        section.orientation = WD_ORIENT.PORTRAIT
        section.page_width = Inches(8.27)
        section.page_height = Inches(11.69)
        section.left_margin = Inches(0.45)
        section.right_margin = Inches(0.45)
        section.top_margin = Inches(0.45)
        section.bottom_margin = Inches(0.45)


def style_document(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = "Malgun Gothic"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    normal.font.size = Pt(10)
    normal.font.color.rgb = RGBColor(26, 25, 22)


def add_title(document: Document) -> None:
    set_section(document.sections[0], "portrait")
    for _ in range(4):
        document.add_paragraph()
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Buildus Care")
    run.bold = True
    run.font.size = Pt(34)
    run.font.color.rgb = RGBColor(26, 107, 90)
    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("웹/앱 화면 캡처 인쇄본")
    run.bold = True
    run.font.size = Pt(24)
    meta = document.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run("생성일: 2026-05-18\n웹 화면과 앱/모바일 화면을 분리하고, 긴 화면은 인쇄용으로 분할했습니다.")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(96, 93, 86)
    document.add_page_break()


def add_section_title(document: Document, title_text: str, subtitle: str, orientation: str) -> None:
    section = document.add_section(WD_SECTION.NEW_PAGE)
    set_section(section, orientation)
    for _ in range(2):
        document.add_paragraph()
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(title_text)
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(26, 25, 22)
    sub = document.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run(subtitle)
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(96, 93, 86)


def add_caption_table(document: Document, title: str, meta: str) -> None:
    table = document.add_table(rows=2, cols=1)
    table.autofit = True
    title_cell = table.cell(0, 0)
    meta_cell = table.cell(1, 0)
    set_cell_shading(title_cell, "E6F0ED")
    set_cell_shading(meta_cell, "FFFFFF")
    set_cell_border(title_cell)
    set_cell_border(meta_cell)
    title_para = title_cell.paragraphs[0]
    title_run = title_para.add_run(title)
    title_run.bold = True
    title_run.font.size = Pt(12)
    title_run.font.color.rgb = RGBColor(26, 107, 90)
    meta_para = meta_cell.paragraphs[0]
    meta_run = meta_para.add_run(meta)
    meta_run.font.size = Pt(8)
    meta_run.font.color.rgb = RGBColor(96, 93, 86)


def add_image_segment(document: Document, segment_path: Path, width_inches: float) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.add_run().add_picture(str(segment_path), width=Inches(width_inches))


def make_segments(source_path: Path, group: str, index: int, width_inches: float, content_height_inches: float) -> list[Path]:
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    with Image.open(source_path) as source:
        source = source.convert("RGB")
        pixels_per_inch = source.width / width_inches
        chunk_height = max(300, int(content_height_inches * pixels_per_inch))
        paths = []
        for part, y in enumerate(range(0, source.height, chunk_height), start=1):
            crop = source.crop((0, y, source.width, min(source.height, y + chunk_height)))
            out_path = ASSET_DIR / f"{group}-{index:02d}-{part:02d}.jpg"
            crop.save(out_path, quality=92, optimize=True)
            paths.append(out_path)
        return paths


def add_capture(document: Document, path: Path, group_label: str, index: int, orientation: str, width_inches: float, content_height_inches: float) -> None:
    segments = make_segments(path, group_label, index, width_inches, content_height_inches)
    with Image.open(path) as source:
        original_size = f"{source.width}x{source.height}"

    for part_index, segment in enumerate(segments, start=1):
        document.add_page_break()
        suffix = f" · {part_index}/{len(segments)}" if len(segments) > 1 else ""
        title = f"{group_label.upper()} {index:02d}. {page_label(path)}{suffix}"
        meta = f"{parse_time(path).strftime('%Y-%m-%d %H:%M:%S')} · 원본 {original_size}"
        add_caption_table(document, title, meta)
        spacer = document.add_paragraph()
        spacer.paragraph_format.space_after = Pt(3)
        add_image_segment(document, segment, width_inches)


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    for old in ASSET_DIR.glob("*.jpg"):
        old.unlink()

    existing = [path for path in RAW_PATHS if path.exists()]
    items = sorted(((classify(path), parse_time(path), path) for path in existing), key=lambda row: row[1])
    web = [path for group, _, path in items if group == "web"]
    app = [path for group, _, path in items if group == "app"]

    document = Document()
    style_document(document)
    add_title(document)

    add_section_title(document, "웹 화면", f"{len(web)}개 캡처 · 데스크톱 브라우저 기준 · A4 가로", "landscape")
    for idx, path in enumerate(web, start=1):
        add_capture(document, path, "web", idx, "landscape", width_inches=10.7, content_height_inches=5.55)

    add_section_title(document, "앱/모바일 화면", f"{len(app)}개 캡처 · 모바일 브라우저/앱 화면 기준 · A4 세로", "portrait")
    for idx, path in enumerate(app, start=1):
        add_capture(document, path, "app", idx, "portrait", width_inches=4.7, content_height_inches=8.55)

    document.save(OUT_DOCX)
    print(OUT_DOCX)
    print(f"images={len(existing)} web={len(web)} app={len(app)}")


if __name__ == "__main__":
    main()
